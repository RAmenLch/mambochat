# backend/services/kb_service.py

import asyncio
import json
import logging
import re
import jieba
from abc import ABC, abstractmethod
from typing import List, Tuple, Optional, Set, Dict
from collections import defaultdict
from io import BytesIO

from fastapi import HTTPException
from langchain_openai import OpenAIEmbeddings
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import func, select

from backend import schemas
from backend.crud import kb_crud, resource_crud, provider_crud, setting_crud
from backend.models import resource_model, kb_model
from backend.schemas import kb as kb_schemas
from backend.schemas.enums import (
    ModelType,
    ResourceType,
    ProviderWorkerType,
    KBFileStatus, ResourceItemType
)
from backend.services.file_service import FileService
from backend.services.stream_manager_service import stream_manager
from backend.config.timezone_config import get_configured_now

SUP_DIM = [384, 768, 1024, 1536, 2560, 3072, 4096]

_KNOWN_TEXT_APPLICATION_TYPES = {
    "application/json", "application/xml", "application/sql",
    "application/javascript", "application/x-sh", "application/x-yaml",
    "application/x-ipynb+json",
}

_DOCUMENT_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def _is_vectorizable_text_type(mime_type: str) -> bool:
    """
    判断 MIME 类型是否为可向量化的文本类型。
    - text/* 前缀通用放行
    - 已知的文本类 application/* 类型放行
    - 支持的文档类型（PDF、Word）放行
    """
    if mime_type.startswith("text/"):
        return True
    if mime_type in _KNOWN_TEXT_APPLICATION_TYPES:
        return True
    if mime_type in _DOCUMENT_MIME_TYPES:
        return True
    return False


logger = logging.getLogger(__name__)


# --- Text Splitters ---

class AbstractTextSplitter(ABC):
    def __init__(self, chunk_size: int, chunk_overlap: int):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    @abstractmethod
    def split_text(self, text: str) -> List[str]:
        pass


class SimpleTextSplitter(AbstractTextSplitter):
    """
    简单的文本切分器，优先按换行符切分，再按字符数切分。
    """

    def split_text(self, text: str) -> List[str]:
        if not text:
            return []

        chunks = []
        start = 0
        text_len = len(text)

        while start < text_len:
            end = start + self.chunk_size

            # 如果不是最后一段，尝试寻找最近的换行符作为切分点
            if end < text_len:
                lookback = text.rfind('\n', start, end)
                if lookback != -1 and lookback > start + (self.chunk_size // 2):
                    end = lookback + 1  # 包含换行符

            chunk = text[start:end]
            if chunk.strip():
                chunks.append(chunk)

            # 计算下一次的起始位置，考虑重叠
            start = end - self.chunk_overlap if end < text_len else end

            # 防止死循环（如果 overlap >= size）
            if start >= end:
                start = end

        return chunks


class SepTextSplitter(AbstractTextSplitter):
    """
    基于特定分隔符的切分器。
    如果切分后的段落超过 chunk_size，则回退到 SimpleTextSplitter 进行二次切分。
    """

    def __init__(self, chunk_size: int, chunk_overlap: int, separator: str):
        super().__init__(chunk_size, chunk_overlap)
        self.separator = separator or "\n\n"
        self._fallback_splitter = SimpleTextSplitter(chunk_size, chunk_overlap)

    def split_text(self, text: str) -> List[str]:
        if not text:
            return []

        raw_chunks = text.split(self.separator)
        final_chunks = []

        for raw in raw_chunks:
            if not raw.strip():
                continue

            # 如果单段长度超过限制，使用 fallback 切分
            if len(raw) > self.chunk_size:
                sub_chunks = self._fallback_splitter.split_text(raw)
                final_chunks.extend(sub_chunks)
            else:
                final_chunks.append(raw)

        return final_chunks


class MarkdownTextSplitter(AbstractTextSplitter):
    """
    基于 Markdown 标题层级的语义切分器。
    按 Markdown 标题（# ~ ######）将文档分成语义段落，保留标题作为上下文前缀。
    如果单个段落超过 chunk_size，则回退到 SimpleTextSplitter 进行二次切分。
    """

    _HEADING_PATTERN = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)

    def split_text(self, text: str) -> List[str]:
        if not text:
            return []

        # 按标题行切分，找到所有标题位置
        lines = text.split('\n')
        chunks: List[List[str]] = []
        current_chunk: List[str] = []
        current_headings: List[str] = []

        for line in lines:
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                title = heading_match.group(2).strip()

                # 遇到标题时，将之前的内容保存为一个 chunk
                if current_chunk:
                    # 将当前标题层级之前的标题作为上下文前缀
                    # 保留同级或更高级别的标题
                    prefix_headings = [h for h in current_headings if len(h[0]) <= level]
                    prefix_headings.append((heading_match.group(1), title))
                    header_prefix = '\n'.join(f"{h[0]} {h[1]}" for h in prefix_headings)

                    chunk_text = '\n'.join(current_chunk).strip()
                    if chunk_text:
                        chunks.append([chunk_text])

                    current_chunk = []
                    current_headings = prefix_headings
                else:
                    # 空 chunk，只更新标题栈
                    current_headings = [h for h in current_headings if len(h[0]) <= level]
                    current_headings.append((heading_match.group(1), title))
            else:
                current_chunk.append(line)

        # 保存最后一段
        if current_chunk:
            chunk_text = '\n'.join(current_chunk).strip()
            if chunk_text:
                chunks.append([chunk_text])

        # 对每个 chunk 添加标题前缀，并检查长度
        fallback_splitter = SimpleTextSplitter(self.chunk_size, self.chunk_overlap)
        final_chunks: List[str] = []
        current_headings_str = ""

        for idx, chunk_lines in enumerate(chunks):
            # 尝试从 chunk 内容中提取最近的标题行作为前缀
            heading_prefix = self._extract_heading_context(chunk_lines[0] if chunk_lines else "")
            if heading_prefix:
                current_headings_str = heading_prefix

            full_text = chunk_lines[0]

            # 如果需要，添加标题上下文前缀
            if current_headings_str:
                full_text = current_headings_str + "\n\n" + full_text

            # 如果单段超过限制，使用 fallback 切分
            if len(full_text) > self.chunk_size:
                sub_chunks = fallback_splitter.split_text(full_text)
                final_chunks.extend(sub_chunks)
            else:
                final_chunks.append(full_text)

        return final_chunks

    @staticmethod
    def _extract_heading_context(text: str) -> str:
        """从文本开头提取连续的 Markdown 标题行作为上下文"""
        lines = text.split('\n')
        heading_lines = []
        for line in lines:
            if re.match(r'^#{1,6}\s+', line):
                heading_lines.append(line)
            elif line.strip():
                break
        return '\n'.join(heading_lines) if heading_lines else ""


class SplitterFactory:
    @staticmethod
    def create(config: kb_schemas.KBTextSplitterConfig) -> AbstractTextSplitter:
        if config.splitter_type == kb_schemas.KBSplitterType.SEPARATOR:
            return SepTextSplitter(
                chunk_size=config.chunk_size,
                chunk_overlap=config.chunk_overlap,
                separator=config.separator
            )
        elif config.splitter_type == kb_schemas.KBSplitterType.MARKDOWN:
            return MarkdownTextSplitter(
                chunk_size=config.chunk_size,
                chunk_overlap=config.chunk_overlap
            )
        else:
            return SimpleTextSplitter(
                chunk_size=config.chunk_size,
                chunk_overlap=config.chunk_overlap
            )


# --- Content Extractors ---

class AbstractContentExtractor(ABC):
    @abstractmethod
    async def extract(self, resource: schemas.ResourceWithVersions, db: AsyncSession) -> str:
        """从资源中提取纯文本内容"""
        pass


class FileExtractor(AbstractContentExtractor):
    """适用于 FILE 和 KB_FILE 类型，从文件服务读取文件内容（纯文本文件）"""

    async def _read_file_bytes(self, resource: schemas.ResourceWithVersions, db: AsyncSession) -> bytes:
        """读取文件原始字节内容"""
        if not resource.latest_version or not resource.latest_version.content:
            raise ValueError("Resource content (file_id) is empty.")

        file_id = resource.latest_version.content
        file_service = FileService(db)

        db_file = await file_service.get_file(file_id)
        if not db_file:
            raise ValueError(f"File record not found for ID: {file_id}")

        if not _is_vectorizable_text_type(db_file.mime_type):
            raise ValueError(
                f"Unsupported file type for vectorization: {db_file.mime_type}.")

        return await file_service.get_file_content(file_id)

    async def extract(self, resource: schemas.ResourceWithVersions, db: AsyncSession) -> str:
        if not resource.latest_version or not resource.latest_version.content:
            raise ValueError("Resource content (file_id) is empty.")

        file_id = resource.latest_version.content
        file_service = FileService(db)

        db_file = await file_service.get_file(file_id)
        if not db_file:
            raise ValueError(f"File record not found for ID: {file_id}")

        if not _is_vectorizable_text_type(db_file.mime_type):
            raise ValueError(
                f"Unsupported file type for vectorization: {db_file.mime_type}. Only text files are supported.")

        try:
            content_bytes = await file_service.get_file_content(file_id)
            try:
                return content_bytes.decode('utf-8')
            except UnicodeDecodeError:
                return content_bytes.decode('latin-1')
        except Exception as e:
            raise ValueError(f"Failed to read file content: {e}")


class PdfExtractor(AbstractContentExtractor):
    """从 PDF 文件提取文本（Markdown 格式），使用 PyMuPDF4LLM"""

    async def extract(self, resource: schemas.ResourceWithVersions, db: AsyncSession) -> str:
        if not resource.latest_version or not resource.latest_version.content:
            raise ValueError("Resource content (file_id) is empty.")

        file_id = resource.latest_version.content
        file_service = FileService(db)

        db_file = await file_service.get_file(file_id)
        if not db_file:
            raise ValueError(f"File record not found for ID: {file_id}")

        if db_file.mime_type != "application/pdf":
            raise ValueError(f"PdfExtractor expects PDF file, got: {db_file.mime_type}")

        try:
            content_bytes = await file_service.get_file_content(file_id)
            from langchain_pymupdf4llm import PyMuPDF4LLMLoader
            from tempfile import NamedTemporaryFile
            import os

            # PyMuPDF4LLMLoader 需要文件路径，写入临时文件
            with NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(content_bytes)
                tmp_path = tmp.name

            try:
                loader = PyMuPDF4LLMLoader(file_path=tmp_path)
                docs = loader.load()
                text = "\n\n".join(doc.page_content for doc in docs)
            finally:
                os.unlink(tmp_path)

            return text if text.strip() else ""
        except Exception as e:
            raise ValueError(f"Failed to extract PDF content: {e}")


class DocxExtractor(AbstractContentExtractor):
    """从 Word (.docx) 文件提取文本，使用 python-docx"""

    async def extract(self, resource: schemas.ResourceWithVersions, db: AsyncSession) -> str:
        if not resource.latest_version or not resource.latest_version.content:
            raise ValueError("Resource content (file_id) is empty.")

        file_id = resource.latest_version.content
        file_service = FileService(db)

        db_file = await file_service.get_file(file_id)
        if not db_file:
            raise ValueError(f"File record not found for ID: {file_id}")

        expected_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if db_file.mime_type != expected_mime:
            raise ValueError(f"DocxExtractor expects DOCX file, got: {db_file.mime_type}")

        try:
            content_bytes = await file_service.get_file_content(file_id)
            import docx

            document = docx.Document(BytesIO(content_bytes))
            text_parts = []
            for para in document.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # 也提取表格内容
            for table in document.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to extract DOCX content: {e}")


class TextExtractor(AbstractContentExtractor):
    """适用于 SYSTEM_PROMPT 和 SUBMESSAGE_TEMPLATE，直接读取版本内容"""

    async def extract(self, resource: schemas.ResourceWithVersions, db: AsyncSession) -> str:
        if not resource.latest_version:
            return ""
        return resource.latest_version.content or ""


class ExtractorFactory:
    @staticmethod
    def get_extractor(resource_type: str, mime_type: str = None) -> AbstractContentExtractor:
        if resource_type in [ResourceType.FILE.value, ResourceType.KB_FILE.value]:
            if mime_type == "application/pdf":
                return PdfExtractor()
            elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return DocxExtractor()
            else:
                return FileExtractor()
        elif resource_type in [ResourceType.SYSTEM_PROMPT.value, ResourceType.SUBMESSAGE_TEMPLATE.value]:
            return TextExtractor()
        else:
            raise ValueError(f"Unsupported resource type for extraction: {resource_type}")


# --- Service ---

class KnowledgeBaseService:
    # 类级别集合，用于记录当前正在运行的任务ID (resource_id)，实现简单的互斥锁
    _running_tasks: Set[str] = set()

    def __init__(self, db: AsyncSession):
        self.db = db

    async def _get_embedding_client(self, model_id: str) -> Tuple[OpenAIEmbeddings, int]:
        """
        根据 model_id 获取配置好的 LangChain Embedding 客户端和维度。
        """
        model = await provider_crud.get_model(self.db, model_id)
        if not model or not model.provider:
            raise ValueError(f"Model {model_id} not found or provider missing.")

        if model.model_type != ModelType.EMBEDDING.value:
            raise ValueError(f"Model {model_id} is not an embedding model.")

        meta_config = json.loads(model.meta_config) if model.meta_config else {}
        dimension = meta_config.get("embedding_dimension")

        if not dimension:
            raise ValueError(f"Model {model_id} configuration is missing 'embedding_dimension'.")

        supported_dims = SUP_DIM
        if dimension not in supported_dims:
            raise ValueError(f"Dimension {dimension} is not supported. Supported: {supported_dims}")

        provider = model.provider

        if provider.worker_type in [ProviderWorkerType.OPENAI.value, ProviderWorkerType.DEEPSEEK.value]:
            client = OpenAIEmbeddings(
                model=model.modelId,
                api_key=provider.apiKey,
                base_url=provider.apiHost,
                check_embedding_ctx_length=False
            )
            return client, dimension
        elif provider.worker_type == ProviderWorkerType.GOOGLE.value:
            raise ValueError("Google Embeddings are not currently supported. Please use an OpenAI-compatible provider.")
        else:
            raise ValueError(f"Unsupported provider worker type for embeddings: {provider.worker_type}")

    async def _validate_kb_hierarchy(self, parent_id: str) -> schemas.Resource:
        """
        验证层级约束：
        1. 资源必须位于 KNOWLEDGE_BASE 类型节点下。
        2. 返回该 KB 节点资源对象。
        """
        if not parent_id:
            raise HTTPException(status_code=400, detail="Parent ID is required for KB resources.")

        # 获取所有祖先节点
        ancestors = await resource_crud.get_batch_resource_ancestors(self.db, [parent_id])

        kb_nodes = [
            res for res in ancestors
            if res.resourceType == ResourceType.KNOWLEDGE_BASE.value
        ]

        if len(kb_nodes) == 0:
            raise HTTPException(status_code=400, detail="Resource must be within a Knowledge Base.")

        if len(kb_nodes) > 1:
            raise HTTPException(status_code=400, detail="Nested Knowledge Bases are not allowed.")

        return kb_nodes[0]

    async def _cleanup_vectors(self, resource_id: str, dimension: int):
        """
        清理指定资源的所有索引数据（向量 + 全文检索）。
        注意：方法名保持为 _cleanup_vectors 以兼容 resource_service 的调用，但逻辑已扩展。
        """
        # 1. 清理向量
        vector_ids = await kb_crud.get_vector_ids_by_resource(self.db, resource_id)
        if vector_ids:
            await kb_crud.delete_vectors(self.db, dimension, vector_ids)

        # 2. 清理 FTS 索引
        fts_ids = await kb_crud.get_fts_ids_by_resource(self.db, resource_id)
        if fts_ids:
            await kb_crud.delete_fts_indexes(self.db, fts_ids)

    async def create_knowledge_base(self, kb_data: kb_schemas.KBCreate) -> schemas.Resource:
        """
        创建知识库资源。
        """
        # 1. 校验模型
        try:
            model = await provider_crud.get_model(self.db, kb_data.embedding_model_id)
            if not model or not model.provider:
                raise ValueError(f"Model {kb_data.embedding_model_id} not found.")

            if model.model_type != ModelType.EMBEDDING.value:
                raise ValueError("Selected model is not an embedding model.")

            meta_config = json.loads(model.meta_config) if model.meta_config else {}
            dimension = meta_config.get("embedding_dimension")

            supported_dims = [384, 768, 1024, 1536, 2560, 3072, 4096]
            if dimension not in supported_dims:
                raise ValueError(f"Model dimension {dimension} is not supported. Supported: {supported_dims}")

        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))

        # 2. 准备 Attributes
        attributes = {
            "embedding_model_id": kb_data.embedding_model_id,
            "dimension": dimension,
            "embedding_rate_limit": kb_data.embedding_rate_limit
        }

        # 3. 创建 Resource
        resource_create = schemas.ResourceCreate(
            name=kb_data.name,
            description=kb_data.description,
            itemType=ResourceItemType.FOLDER,
            resourceType=ResourceType.KNOWLEDGE_BASE,
            parentId=kb_data.parent_id,
            initial_attributes=attributes,
            initial_content=""
        )

        new_resource = await resource_crud.create_resource(self.db, resource_create)

        # 4. 手动创建初始版本并关联
        initial_version = resource_model.ResourceVersion(
            resourceId=new_resource.id,
            name="初始配置",
            content="",
            attributes=attributes
        )
        self.db.add(initial_version)
        await self.db.flush()

        new_resource.latestVersionId = initial_version.id
        await self.db.commit()
        await self.db.refresh(new_resource)
        await self.db.refresh(new_resource, ['latest_version'])

        return new_resource

    async def update_kb_file_config(self, resource_id: str,
                                    config_data: kb_schemas.KBUpdateConfigRequest) -> schemas.Resource:
        """
        更新知识库文件的切分配置。
        配置现在存储在 Resource.kb_config 中，所有版本共享。
        """
        resource = await resource_crud.get_resource(self.db, resource_id)
        if not resource:
            raise HTTPException(status_code=404, detail="Resource not found")

        # 更新 Resource 表字段
        resource.kb_config = config_data.splitter_config.model_dump()

        await self.db.commit()
        await self.db.refresh(resource)
        return resource

    async def delete_kb_file(self, resource_id: str) -> schemas.Resource:
        """
        删除 KB 文件资源：先清理向量和FTS数据，再删除资源记录。
        """
        resource = await resource_crud.get_resource(self.db, resource_id)
        if not resource:
            raise HTTPException(status_code=404, detail="Resource not found")

        try:
            if resource.parentId:
                kb_resource = await self._validate_kb_hierarchy(resource.parentId)
                if kb_resource and kb_resource.latest_version and kb_resource.latest_version.attributes:
                    dimension = kb_resource.latest_version.attributes.get("dimension")
                    if dimension:
                        await self._cleanup_vectors(resource_id, dimension)
        except Exception:
            pass

        await kb_crud.delete_chunks_by_resource(self.db, resource_id)
        return await resource_crud.delete_resource(self.db, resource_id)

    async def get_comprehensive_file_status(self, resource_id: str) -> kb_schemas.KBProcessingStatus:
        """
        获取文件的综合处理状态，包含过时判定 (Staleness)。
        """
        # 1. 获取数据库统计信息
        stats = await kb_crud.get_chunk_stats_by_resource(self.db, resource_id)

        # 2. 检查内存任务状态
        is_running = resource_id in KnowledgeBaseService._running_tasks

        # 3. 状态修正
        if is_running:
            stats.file_status = KBFileStatus.EMBEDDING
        elif stats.pending_chunks > 0:
            stats.file_status = KBFileStatus.FAILED
            stats.failed_chunks += stats.pending_chunks
            stats.pending_chunks = 0

        # 4. 过时判定 (Staleness Check)
        # 获取资源最新版本时间
        resource = await resource_crud.get_resource(self.db, resource_id)
        if resource and resource.latest_version:
            latest_version_time = resource.latest_version.updatedAt

            # 获取最新的 Chunk 处理时间
            stmt = select(func.max(kb_model.ResourceKBChunk.processed_at)).where(
                kb_model.ResourceKBChunk.resource_id == resource_id
            )
            result = await self.db.execute(stmt)
            max_processed_at = result.scalar()

            # 如果版本更新时间晚于最后一次向量处理时间，标记为过时
            if max_processed_at and latest_version_time > max_processed_at:
                stats.is_stale = True
            elif not max_processed_at and stats.total_chunks == 0:
                pass

        return stats

    async def _publish_status(self, resource_id: str, status: KBFileStatus,
                              total: int, completed: int, failed: int, stopped: int,
                              error_message: str = None):
        """
        辅助方法：构建统一的 KBProcessingStatus 并推送。
        """
        pending = max(0, total - completed - failed - stopped)

        data = kb_schemas.KBProcessingStatus(
            resource_id=resource_id,
            total_chunks=total,
            pending_chunks=pending,
            completed_chunks=completed,
            failed_chunks=failed,
            stopped_chunks=stopped,
            file_status=status,
            error_message=error_message
        )
        await stream_manager.publish(resource_id, data.model_dump())

    async def handle_task_action(self, resource_id: str, request: kb_schemas.KBRunTaskRequest):
        """
        处理任务控制动作：Start, Resume, Stop。
        """
        # 1. 停止任务
        if request.action == kb_schemas.KBTaskAction.STOP:
            if resource_id in KnowledgeBaseService._running_tasks:
                await stream_manager.request_cancellation(resource_id)
                await kb_crud.mark_pending_chunks_as_stopped(self.db, resource_id)
                return {"message": "Cancellation requested."}
            else:
                return {"message": "Task is not running."}

        # 2. 检查任务是否已在运行
        if resource_id in KnowledgeBaseService._running_tasks:
            raise HTTPException(status_code=409, detail="Task is already running for this file.")

        # 3. 获取资源和父 KB 信息
        resource = await resource_crud.get_resource_with_versions(self.db, resource_id)
        if not resource:
            raise HTTPException(status_code=404, detail="Resource not found")

        kb_resource = await self._validate_kb_hierarchy(resource.parentId)
        kb_attrs = kb_resource.latest_version.attributes or {}
        embedding_model_id = kb_attrs.get("embedding_model_id")
        dimension = kb_attrs.get("dimension")
        rate_limit = kb_attrs.get("embedding_rate_limit", 0.0)
        kb_id = kb_resource.id

        if not embedding_model_id or not dimension:
            raise HTTPException(status_code=400, detail="Knowledge Base configuration is incomplete.")

        # 4. 获取切分配置 (从 Resource 层级)
        kb_config_dict = resource.kb_config

        # 如果配置为空且动作是 START，则尝试加载全局默认配置
        if not kb_config_dict and request.action == kb_schemas.KBTaskAction.START:
            # 获取全局默认切分参数
            size_setting = await setting_crud.get_setting(self.db, "kb_default_chunk_size")
            overlap_setting = await setting_crud.get_setting(self.db, "kb_default_chunk_overlap")

            # 解析参数，如果解析失败或未设置则使用硬编码默认值
            try:
                default_size = int(size_setting.value) if size_setting and size_setting.value else 500
            except (ValueError, TypeError):
                default_size = 500

            try:
                default_overlap = int(overlap_setting.value) if overlap_setting and overlap_setting.value else 50
            except (ValueError, TypeError):
                default_overlap = 50

            # 构建默认配置 (默认使用 simple 切分器)
            kb_config_dict = {
                "splitter_type": "simple",
                "chunk_size": default_size,
                "chunk_overlap": default_overlap
            }

            # 将默认配置保存回资源，以便下次使用
            resource.kb_config = kb_config_dict
            await self.db.commit()
            await self.db.refresh(resource)

        # 获取上次运行的配置快照 (用于 Resume 校验)
        current_attributes = dict(resource.latest_version.attributes) if resource.latest_version.attributes else {}
        last_config_dict = current_attributes.get("last_ingest_config")

        target_splitter_config = None

        # 5. 动作分发
        if request.action == kb_schemas.KBTaskAction.RESUME:
            if not last_config_dict:
                raise HTTPException(status_code=400, detail="No previous task found. Please use START.")

            # 检查配置一致性
            if kb_config_dict != last_config_dict:
                raise HTTPException(
                    status_code=409,
                    detail={
                        "message": "Configuration mismatch. Cannot resume task.",
                        "current_config": kb_config_dict,
                        "last_ingest_config": last_config_dict
                    }
                )

            target_splitter_config = kb_schemas.KBTextSplitterConfig(**last_config_dict)

            asyncio.create_task(self._run_embedding_loop(
                resource_id, embedding_model_id, dimension, rate_limit,
                resume=True,
                splitter_config=target_splitter_config,
                kb_id=kb_id
            ))

        elif request.action == kb_schemas.KBTaskAction.START:
            # 更新 last_ingest_config 到当前版本
            current_attributes["last_ingest_config"] = kb_config_dict

            await resource_crud.update_resource_version(
                self.db,
                resource.latestVersionId,
                schemas.ResourceVersionUpdate(attributes=current_attributes)
            )

            target_splitter_config = kb_schemas.KBTextSplitterConfig(**kb_config_dict)

            asyncio.create_task(self._run_embedding_loop(
                resource_id, embedding_model_id, dimension, rate_limit,
                resume=False,
                splitter_config=target_splitter_config,
                kb_id=kb_id
            ))

        return {"message": "Task started."}

    async def _run_embedding_loop(
            self,
            resource_id: str,
            model_id: str,
            dimension: int,
            rate_limit: float,
            resume: bool,
            splitter_config: kb_schemas.KBTextSplitterConfig,
            kb_id: str = ""
    ):
        """
        核心任务循环：处理切分、嵌入、存储(向量+FTS)、状态更新和取消。
        """
        KnowledgeBaseService._running_tasks.add(resource_id)

        total_count = 0
        processed_count = 0
        failed_count = 0
        stopped_count = 0

        from backend.database import AsyncSessionLocal
        async with AsyncSessionLocal() as session:
            try:
                temp_service = KnowledgeBaseService(session)
                try:
                    client, _ = await temp_service._get_embedding_client(model_id)
                except Exception as e:
                    logger.error(f"Model init failed for resource {resource_id}: {e}", exc_info=True)
                    await self._publish_status(resource_id, KBFileStatus.FAILED, 0, 0, 0, 0,
                                               error_message=f"Embedding模型初始化失败: {e}")
                    return

                # --- 阶段 1: 准备数据 (Start 模式需切分) ---
                if not resume:
                    # 1. 获取资源
                    resource = await resource_crud.get_resource_with_versions(session, resource_id)
                    if not resource:
                        raise ValueError("Resource not found.")

                    # 2. 清理旧数据 (向量 + FTS + Chunks)
                    await self._publish_status(resource_id, KBFileStatus.CLEANING, 0, 0, 0, 0)
                    await temp_service._cleanup_vectors(resource_id, dimension)
                    await kb_crud.delete_chunks_by_resource(session, resource_id)

                    # 3. 提取文本
                    await self._publish_status(resource_id, KBFileStatus.READING, 0, 0, 0, 0)

                    try:
                        # 获取文件的 MIME 类型以选择正确的提取器
                        file_mime_type = None
                        if resource.resourceType in [ResourceType.FILE.value, ResourceType.KB_FILE.value]:
                            if resource.latest_version and resource.latest_version.content:
                                file_svc = FileService(session)
                                db_file = await file_svc.get_file(resource.latest_version.content)
                                if db_file:
                                    file_mime_type = db_file.mime_type

                        extractor = ExtractorFactory.get_extractor(resource.resourceType, mime_type=file_mime_type)
                        text_content = await extractor.extract(resource, session)
                    except Exception as e:
                        logger.error(f"Extraction failed for resource {resource_id}: {e}", exc_info=True)
                        raise ValueError(f"Content extraction failed: {e}")

                    # 4. 切分
                    await self._publish_status(resource_id, KBFileStatus.SPLITTING, 0, 0, 0, 0)
                    splitter = SplitterFactory.create(splitter_config)
                    text_chunks = splitter.split_text(text_content)

                    # 5. 存储 Chunks (包含 created_at)
                    chunk_schemas = [
                        kb_schemas.KBChunkCreate(
                            resource_id=resource_id,
                            content=chunk,
                            chunk_index=idx,
                            byte_size=len(chunk.encode('utf-8'))
                        )
                        for idx, chunk in enumerate(text_chunks)
                    ]
                    await kb_crud.batch_create_chunks(session, chunk_schemas)

                    total_count = len(chunk_schemas)
                else:
                    # Resume 模式
                    stats = await kb_crud.get_chunk_stats_by_resource(session, resource_id)
                    total_count = stats.total_chunks
                    processed_count = stats.completed_chunks
                    failed_count = 0
                    stopped_count = 0

                # --- 阶段 2: 嵌入循环 ---

                target_statuses = [kb_schemas.KBChunkStatus.PENDING.value]
                if resume:
                    target_statuses.append(kb_schemas.KBChunkStatus.FAILED.value)
                    target_statuses.append(kb_schemas.KBChunkStatus.STOPPED.value)

                pending_chunks = await kb_crud.get_chunks_by_statuses(session, resource_id, target_statuses)

                if resume:
                    stats = await kb_crud.get_chunk_stats_by_resource(session, resource_id)
                    total_count = stats.total_chunks
                    processed_count = stats.completed_chunks

                batch_size = 10

                await self._publish_status(resource_id, KBFileStatus.EMBEDDING,
                                           total_count, processed_count, failed_count, stopped_count)

                for i in range(0, len(pending_chunks), batch_size):
                    if await stream_manager.is_cancellation_requested(resource_id):
                        remaining = len(pending_chunks) - i
                        stopped_count += remaining
                        await kb_crud.mark_pending_chunks_as_stopped(session, resource_id)
                        await self._publish_status(resource_id, KBFileStatus.STOPPED,
                                                   total_count, processed_count, failed_count, stopped_count)
                        return

                    batch = pending_chunks[i:i + batch_size]
                    texts = [c.content for c in batch]

                    try:
                        # 1. 生成向量
                        vectors = await client.aembed_documents(texts)

                        current_batch_success = 0
                        current_batch_failed = 0
                        now = get_configured_now()

                        for idx, vector in enumerate(vectors):
                            chunk = batch[idx]
                            vec_rowid = None
                            fts_rowid = None

                            try:
                                # 2. 插入向量
                                if len(vector) == dimension:
                                    vec_rowid = await kb_crud.insert_vector(
                                        session, dimension, vector,
                                        kb_id=kb_id, resource_id=resource_id
                                    )

                                # 3. 插入 FTS 索引 (使用 jieba 分词)
                                tokens = " ".join(jieba.cut_for_search(chunk.content))
                                fts_rowid = await kb_crud.insert_fts_index(session, tokens)

                                # 4. 验证双路插入结果
                                if vec_rowid is not None and fts_rowid is not None:
                                    # 更新 Chunk 状态
                                    chunk.vector_id = vec_rowid
                                    chunk.fts_id = fts_rowid
                                    chunk.status = kb_schemas.KBChunkStatus.COMPLETED.value
                                    chunk.processed_at = now
                                    session.add(chunk)
                                    current_batch_success += 1
                                else:
                                    # 任何一路失败，回滚该条目
                                    raise Exception("Dual index insertion incomplete")

                            except Exception as inner_e:
                                # 记录详细的错误信息
                                error_str = f"[Chunk #{chunk.chunk_index}] {type(inner_e).__name__}: {str(inner_e)}"
                                logger.error(f"Chunk embedding failed for resource {resource_id}: {error_str}")
                                # 回滚补偿：如果部分插入成功，尝试清理
                                if vec_rowid:
                                    await kb_crud.delete_vectors(session, dimension, [vec_rowid])
                                if fts_rowid:
                                    await kb_crud.delete_fts_indexes(session, [fts_rowid])

                                chunk.vector_id = None
                                chunk.fts_id = None
                                chunk.status = kb_schemas.KBChunkStatus.FAILED.value
                                chunk.error_message = error_str[:500]  # 截断防止过长
                                session.add(chunk)
                                current_batch_failed += 1

                        await session.commit()

                        # 处理向量生成数量少于文本数量的情况 (极少发生)
                        if len(vectors) < len(batch):
                            diff = len(batch) - len(vectors)
                            current_batch_failed += diff
                            missing_msg = f"Embedding API returned fewer vectors ({len(vectors)}) than texts ({len(batch)})"
                            logger.error(f"Batch embedding partial failure for resource {resource_id}: {missing_msg}")
                            for k in range(len(vectors), len(batch)):
                                chunk = batch[k]
                                chunk.vector_id = None
                                chunk.fts_id = None
                                chunk.status = kb_schemas.KBChunkStatus.FAILED.value
                                chunk.error_message = f"[Chunk #{chunk.chunk_index}] {missing_msg}"
                                session.add(chunk)
                            await session.commit()

                        processed_count += current_batch_success
                        failed_count += current_batch_failed

                    except Exception as e:
                        logger.error(f"Embedding batch failed for resource {resource_id}: {e}", exc_info=True)
                        batch_error = f"Batch embedding API call failed: {type(e).__name__}: {str(e)}"
                        for chunk in batch:
                            chunk.vector_id = None
                            chunk.fts_id = None
                            chunk.status = kb_schemas.KBChunkStatus.FAILED.value
                            chunk.error_message = batch_error[:500]
                            session.add(chunk)
                        await session.commit()
                        failed_count += len(batch)

                    await self._publish_status(resource_id, KBFileStatus.EMBEDDING,
                                               total_count, processed_count, failed_count, stopped_count)

                    if rate_limit > 0:
                        await asyncio.sleep(rate_limit)

                await self._publish_status(resource_id, KBFileStatus.COMPLETED,
                                           total_count, processed_count, failed_count, stopped_count)

            except Exception as e:
                error_msg = f"{type(e).__name__}: {str(e)}"
                logger.error(f"Task failed for resource {resource_id}: {error_msg}", exc_info=True)
                failed_count = total_count - processed_count - stopped_count
                await self._publish_status(resource_id, KBFileStatus.FAILED,
                                           total_count, processed_count, failed_count, stopped_count,
                                           error_message=error_msg)
            finally:
                KnowledgeBaseService._running_tasks.discard(resource_id)
                await stream_manager.close_stream(resource_id)

    async def search_kb(self, request: kb_schemas.KBSearchRequest) -> kb_schemas.KBSearchResponse:
        """
        执行混合检索 (Vector + BM25) 并使用 RRF 融合排序。
        支持 resource_name 解析和 chunk_index 范围后过滤：
        不在 vec 表使用 MATCH 按 chunk_index 过滤，而是先做向量距离搜索扩大候选集，
        再仅对匹配 chunk_index 范围的结果进行筛选。
        """
        target_kb_id = request.kb_id
        embedding_model_id = None
        target_resource_id = request.resource_id

        if target_kb_id:
            kb_resource = await resource_crud.get_resource(self.db, target_kb_id)
            if kb_resource and kb_resource.latest_version:
                embedding_model_id = (kb_resource.latest_version.attributes or {}).get("embedding_model_id")

            # 解析 resource_name -> resource_id
            if request.resource_name and not target_resource_id:
                named_resource = await resource_crud.get_resource_by_name_and_parent(
                    self.db, request.resource_name, target_kb_id
                )
                if named_resource:
                    target_resource_id = named_resource.id

        if not embedding_model_id:
            raise HTTPException(status_code=400,
                                detail="Embedding model not determined. Please specify a valid Knowledge Base ID.")

        try:
            client, dimension = await self._get_embedding_client(embedding_model_id)
        except ValueError as e:
            raise HTTPException(status_code=500, detail=str(e))

        # 是否需要 chunk_index 后过滤
        need_index_filter = request.index_start is not None or request.index_end is not None

        # 1. 扩大候选集 (Candidate Multiplier)
        # 当无 chunk_index 过滤时使用 MATCH KNN + 后过滤策略，需扩大候选集
        candidate_multiplier = 2
        candidate_k = request.top_k * candidate_multiplier

        # 2. 并行执行双路召回
        # Vector Path
        async def _vector_search():
            try:
                query_vector = await client.aembed_query(request.query_text)
                if need_index_filter and target_resource_id:
                    # 小数据集场景：JOIN vec + chunk 表，暴力计算距离 + chunk_index 过滤
                    # 不使用 MATCH KNN，避免因后过滤丢掉近距离结果
                    return await kb_crud.search_vectors_with_chunk_filter(
                        self.db, dimension, query_vector, request.top_k,
                        kb_id=target_kb_id,
                        resource_id=target_resource_id,
                        index_start=request.index_start,
                        index_end=request.index_end
                    )
                else:
                    # 常规场景：MATCH KNN + 分区键预过滤
                    return await kb_crud.search_vectors(
                        self.db, dimension, query_vector, candidate_k,
                        kb_id=target_kb_id,
                        resource_id=target_resource_id
                    )
            except Exception as e:
                logger.error(f"Vector search failed: {e}")
                return []

        # Keyword Path (BM25)
        async def _keyword_search():
            try:
                # 使用 jieba 对 Query 分词，并用 OR 连接以提高召回
                keywords = list(jieba.cut(request.query_text))
                if not keywords:
                    return []
                # 构造 MATCH 查询: "term1 OR term2"
                # 注意处理关键词中的特殊字符以防 SQL 注入或语法错误 (简单处理：去除双引号)
                sanitized_keywords = [k.replace('"', '') for k in keywords if k.strip()]
                if not sanitized_keywords:
                    return []
                match_query = " OR ".join(f'"{k}"' for k in sanitized_keywords)
                return await kb_crud.search_fts_bm25(self.db, match_query, candidate_k)
            except Exception as e:
                logger.error(f"Keyword search failed: {e}")
                return []

        vec_results, fts_results = await asyncio.gather(_vector_search(), _keyword_search())

        # 3. 获取所有涉及的 Chunk 信息 (带 chunk_index 范围后过滤)
        vec_ids = [r[0] for r in vec_results]
        fts_ids = [r[0] for r in fts_results]

        if not vec_ids and not fts_ids:
            return kb_schemas.KBSearchResponse(total=0, items=[])

        chunk_rows = await kb_crud.get_chunks_by_mixed_ids(
            self.db, vec_ids, fts_ids, kb_id_filter=target_kb_id,
            index_start=request.index_start,
            index_end=request.index_end
        )

        # 建立映射: Chunk ID -> Chunk Object & Metadata
        chunk_map = {}
        # 建立映射: Vector/FTS ID -> Chunk ID (用于 RRF 评分归属)
        vec_id_to_chunk_id = {}
        fts_id_to_chunk_id = {}

        for row in chunk_rows:
            chunk = row[0]
            chunk_map[chunk.id] = row
            if chunk.vector_id:
                vec_id_to_chunk_id[chunk.vector_id] = chunk.id
            if chunk.fts_id:
                fts_id_to_chunk_id[chunk.fts_id] = chunk.id

        # 4. RRF 融合排序
        # Score = 1 / (k + rank_vec) + 1 / (k + rank_bm25)
        rrf_k = 60
        scores: Dict[str, float] = defaultdict(float)

        # 处理向量结果 (Rank 从 1 开始)
        for rank, (rowid, _) in enumerate(vec_results, start=1):
            c_id = vec_id_to_chunk_id.get(rowid)
            if c_id:
                scores[c_id] += 1.0 / (rrf_k + rank)

        # 处理关键词结果
        for rank, (rowid, _) in enumerate(fts_results, start=1):
            c_id = fts_id_to_chunk_id.get(rowid)
            if c_id:
                scores[c_id] += 1.0 / (rrf_k + rank)

        # 5. 排序并截取 Top K
        sorted_chunk_ids = sorted(scores.keys(), key=lambda x: scores[x], reverse=True)[:request.top_k]

        # 6. 组装最终结果
        items = []
        for c_id in sorted_chunk_ids:
            row = chunk_map.get(c_id)
            if not row:
                continue

            chunk = row[0]
            file_id = row[1]
            file_name = row[2]
            kb_id = row[3]
            kb_name = row[4]
            chunk_index = row[5]

            items.append(kb_schemas.KBSearchResultItem(
                chunk_id=chunk.id,
                chunk_content=chunk.content,
                score=scores[c_id],  # 这里返回 RRF 分数
                resource_id=file_id,
                resource_name=file_name,
                kb_id=kb_id,
                kb_name=kb_name,
                chunk_index=chunk_index
            ))

        return kb_schemas.KBSearchResponse(total=len(items), items=items)
